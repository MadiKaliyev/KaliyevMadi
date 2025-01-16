import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()
faile = doc.add_paragraph('«Hello Wordl!»','Title')
faile.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


par = doc.add_paragraph('')

par.add_run('\tДля вас мы собрали самые известные '
         'и проверенные временем сказки для детей.'
         'Здесь размещены русские народные сказки и '
         'авторские сказки, которые точно стоит прочитать'
         ' ребенку. Детские сказки этого раздела подходят абсолютно'
         ' всем ребятам: подобраны сказки для самых маленьких и для '
         'школьников. Некоторые произведения Вы найдете только у нас,'
         ' в оригинальном изложении!')

doc.save('ворд.docx')
print(faile.text,par.text)